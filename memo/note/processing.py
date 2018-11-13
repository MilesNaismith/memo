from docx import Document
from docx.shared import Inches, Cm, Pt
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shutil import copyfile
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
def get_note(company_name, budget_item, pay_reason, pay_check, pay_sum):
    tomorrow = datetime.date.today() + datetime.timedelta(days=1)
    document = Document()
    ###Шапка###
    p = document.add_paragraph('СОГЛАСОВАНО \nГенеральному директору \nАО «Газпромбанк Лизинг» \nМ.А. Агаджанову') 
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph('')
    p.add_run('Служебная записка на согласование \nадминистративно-хозяйственных расходов от {}'.format(datetime.date.today().strftime('%d.%m.%Y'))).italic = True
    document.add_paragraph('')
    p = document.add_paragraph('Уважаемый Максим Анатольевич!')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('Прошу Вас согласовать расход на административно-хозяйственные нужды:')

    ###Тело###

    ### Название компании
    company = document.add_paragraph(
        'Получатель/контрагент ', style='List Bullet'
    )
    company.add_run(company_name).bold = True

    ### Сумма оплаты
    price = document.add_paragraph(
        'Сумма расхода ', style='List Bullet'
    )
    price.add_run('{} RUB'.format(pay_sum)).bold = True

    ### Дата оплаты
    pay_date = document.add_paragraph(
        'Ожидаемый срок расхода ', style='List Bullet'
    )
    pay_date.add_run(tomorrow.strftime('%d.%m.%Y')).bold = True

    ### Содержимое счета
    pay_explanation = document.add_paragraph(
        'Обоснование расхода', style='List Bullet'
    )
    pay_explanation_text = document.add_paragraph('')
    pay_explanation_text.add_run(pay_reason).bold = True

    ### Номер счета
    pay_doc = document.add_paragraph(
        'Документ, на основании которого осуществляется расход (приложить счет/договор, если применимо) ', style='List Bullet'
    )
    pay_doc.add_run('Счет № {}'.format(pay_check)).bold = True

    ### Статья бюджета
    budget = document.add_paragraph(
        'Статья бюджета* ', style='List Bullet'
    )
    budget.add_run('ИТ расходы/{}'.format(budget_item)).bold = True

    ###Подвал###
    document.add_paragraph('')
    document.add_paragraph('Инициатор расхода                                                                                                                Трусков А.А.')
    document.add_paragraph('Руководитель подразделения инициатора                                                               Карякин М.Б.')   
    document.add_paragraph('Руководитель подразделения ЦФО                                                                             Богданов Д.Т.')
    document.add_paragraph('')
    document.add_paragraph('СОГЛАСОВАНО:')
    document.add_paragraph('Главный бухгалтер                                                                                                            Клочкова Е.В.')
    document.add_paragraph('Финансовый директор                                                                                                    Ястребова Н.Н.')
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    table.height = Cm(0.8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Информация о бюджете'
    hdr_cells[0].width = Cm(5.0)
    hdr_cells[1].width = Cm(12.0)
    document.add_paragraph('')
    lower_colon = document.add_paragraph('')
    run = lower_colon.add_run('* в соответствии с утвержденным Советом директоров на соответствующий год Бюджетом, размещаемым Финансовый департаментом для общего доступа на сервере по адресу: \int.gpbl.ru\Сетевые ресурсы\Справочные материалы\Аналитические отчеты ФД\Бюджет. До утверждения новой редакции Положения о бюджетировании статья бюджета определяется в соответствии с Приложением №2 к Порядку осуществления и контроля административно-хозяйственных расходов ГПБЛ')
    run.italic = True
    run.font.size = Pt(8)
    #ont = run.font
    #font.size = Pt(8)
    print(dir(run))


    document.save('note.docx')
    file_path = os.path.join(BASE_DIR, 'note.docx')
    file_copy_path = os.path.join(BASE_DIR, 'note/media/note.docx')
    copyfile(file_path, file_copy_path)
    
    
if __name__ == "__main__":   
    get_note()    